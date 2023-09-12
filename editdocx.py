from docx import Document
from docx.shared import Inches
import copy

class EditDocx:
    
    def __init__(self, file_path):
        self.orig_document = Document(file_path)
        self.edited_document = copy.deepcopy(self.orig_document)
        self.parsed_orig_document = EditDocx._parse_document(self.orig_document)
        self.parsed_edited_document = EditDocx._parse_document(self.edited_document)
       
    
    def print_parsed_document(self, orig=True):
        if orig:
            self.parsed_orig_document = EditDocx._parse_document(self.orig_document)
            EditDocx._print_dict(self.parsed_orig_document)
        else:
            self.parsed_edited_document = EditDocx._parse_document(self.edited_document)
            EditDocx._print_dict(self.parsed_edited_document)
    
    #<w14:checkbox><w14:checked w14:val="0"/><w14:checkedState w14:val="2612" w14:font="MS Gothic"/><w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/></w14:checkbox> 
    
    def toggle_checkbox(self, para_index, checkbox_index):
        # Step 1: Locate the Paragraph
        paragraph = self.edited_document.paragraphs[para_index]._element
        
        # Step 2: Find the Checkbox
        checkboxes = paragraph.xpath('.//w14:checkbox')
        if checkbox_index < len(checkboxes):
            target_checkbox = checkboxes[checkbox_index]
            
            # Step 3: Toggle Status
            checked_element = target_checkbox.find('.//w14:checked', namespaces={'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'})
            current_val = checked_element.get('{http://schemas.microsoft.com/office/word/2010/wordml}val')
            
            # Toggle the value
            new_val = "1" if current_val == "0" else "0"
            checked_element.set('{http://schemas.microsoft.com/office/word/2010/wordml}val', new_val)
            
            # Locate the common parent <w:sdt> and find the <w:sdtContent> within it
            common_parent = target_checkbox.getparent().getparent()
            sdt_content = common_parent.find('.//w:sdtContent', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if sdt_content is not None:  
                run = sdt_content.find('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                text = run.find('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                text.text = "☒" if new_val == "1" else "☐"
            else:
                print("sdt_content not found.")
            
        else:
            print("Checkbox index out of range.")

    
    
    def insert_paragraph(): #will change # of paragraphs and thus affect subsequent document changes
        pass
    
    def insert_image(self, img_path_or_stream, paragraph_num, height_in_inches): #TODO: allow for width and/or height or NONE
        paragraph = self.edited_document.paragraphs[paragraph_num]
        paragraph.add_run().add_picture(img_path_or_stream, height=Inches(height_in_inches))
    
    def replace_paragraph_text(self, paragraph_num, text):
        pass
        
    def replace_cell_text_in_table(self, table_num, row_num, col_num, text):
        self.edited_document.tables[table_num].cell(row_num, col_num).text = text
        
    def save_edited_docx(self, save_path):
        self.edited_document.save(save_path)
    
    @staticmethod
    def _parse_document(docx_document):
        parsed_document = {}
        parsed_document['Document'] = {'Paragraphs':{}, 'Tables':{}}
        parsed_document['Document']['Paragraphs'] = EditDocx._parse_paragraphs(docx_document.paragraphs)
        parsed_document['Document']['Tables'] = EditDocx._parse_tables(docx_document.tables)
        for s_num, section in enumerate(docx_document.sections):
            parsed_document[f'Section_{s_num}'] = {'Header':{'Paragraphs':{}, 'Tables':{}}, 'Footer':{'Paragraphs':{}, 'Tables':{}}}
            
            if section.header:
                parsed_document[f'Section_{s_num}']['Header']['Paragraphs'] = EditDocx._parse_paragraphs(section.header.paragraphs)
                parsed_document[f'Section_{s_num}']['Header']['Tables']= EditDocx._parse_tables(section.header.tables)
            
            if section.footer:
                parsed_document[f'Section_{s_num}']['Footer']['Paragraphs'] = EditDocx._parse_paragraphs(section.footer.paragraphs)
                parsed_document[f'Section_{s_num}']['Footer']['Tables']= EditDocx._parse_tables(section.footer.tables)
        return parsed_document
    
    @staticmethod
    def _parse_tables(docx_tables):
        parsed_tables = {}
        for t_num, table in enumerate(docx_tables):
            parsed_table = EditDocx._parse_table(table)
            parsed_tables[f'Table_{t_num}'] = parsed_table
            return parsed_tables

    @staticmethod
    def _parse_table(docx_table):
        parsed_table = {}
        for r_num, row in enumerate(docx_table.rows):
            row_data = {}
            for c_num, cell in enumerate(row.cells):
                row_data[f'Col_{c_num}'] = cell.text
            parsed_table[f'Row_{r_num}'] = row_data
        return parsed_table
    
    @staticmethod
    def _parse_paragraphs(docx_paragraphs):
        parsed_paragraphs = {}
        for p_num, paragraph in enumerate(docx_paragraphs):
            parsed_paragraphs[f'Paragraph_{p_num}'] = paragraph.text
        return parsed_paragraphs
    
    @staticmethod
    def _print_dict(d, indent=0):
        for key, value in d.items():
            print('  ' * indent + str(key), end=": ")
            if isinstance(value, dict):
                print()
                EditDocx._print_dict(value, indent + 1)
            elif isinstance(value, list):
                print()
                for i, item in enumerate(value):
                    print('  ' * (indent + 1) + f"[{i}]", end=": ")
                    if isinstance(item, dict):
                        print()
                        EditDocx._print_dict(item, indent + 2)
                    else:
                        print(item)
            else:
                print(value)
                
    @staticmethod
    def _print_docx_document_xml(document, paragraphs_with_checkboxes):
        for paragraph in document.paragraphs:
            p = paragraph._element
            if paragraphs_with_checkboxes:
                checkBoxes = p.xpath('.//w14:checkbox')
                if checkBoxes:
                    print(p.xml)
                    break
            else:
                print(p.xml)
        